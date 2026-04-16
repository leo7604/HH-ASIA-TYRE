"""
Create Customer Getting Started Guide in DOCX format
with screenshots and comprehensive instructions
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_customer_guide():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Configure margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ========== COVER PAGE ==========
    # Add spacing
    for _ in range(4):
        doc.add_paragraph('')
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('🚗 HH ASIA TYRE')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Customer Getting Started Guide')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph('')
    
    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Complete Step-by-Step Guide to Booking Your Appointment Online')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True
    
    for _ in range(3):
        doc.add_paragraph('')
    
    # Website
    website = doc.add_paragraph()
    website.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = website.add_run('https://hh-asia-tyre.vercel.app')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    
    doc.add_paragraph('')
    
    # Version info
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version 1.0 | April 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # Page break
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    toc_title = doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        ('1.', 'Welcome to HH Asia Tyre', '3'),
        ('2.', 'Quick Start - Book in 3 Minutes', '4'),
        ('3.', 'Step 1: Choose Your Branch', '5'),
        ('4.', 'Step 2: Enter Vehicle Details', '7'),
        ('5.', 'Step 3: Select Your Services', '8'),
        ('6.', 'Step 4: Pick Date & Time', '10'),
        ('7.', 'Step 5: Enter Your Information', '12'),
        ('8.', 'Step 6: Review & Confirm', '14'),
        ('9.', 'Our Branch Locations', '15'),
        ('10.', 'Services We Offer', '17'),
        ('11.', 'Understanding Booking Status', '18'),
        ('12.', 'Frequently Asked Questions', '19'),
        ('13.', 'Need Help?', '21'),
    ]
    
    for num, item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num} {item}')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    doc.add_page_break()
    
    # ========== SECTION 1: WELCOME ==========
    doc.add_heading('1. Welcome to HH Asia Tyre', level=1)
    
    doc.add_paragraph(
        'HH Asia Tyre is your trusted automotive service partner with 6 branches across Metro Manila, '
        'Cavite, and Ilocos Norte. We specialize in tire services, oil changes, brake inspection, '
        'battery replacement, and general auto repair.'
    )
    
    doc.add_heading('What You Can Do Online:', level=2)
    
    features = [
        'Book appointments at any of our 6 branches',
        'Choose your preferred date and time',
        'Select specific services you need',
        'Get instant confirmation via email/SMS',
        'Track your booking status'
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    # Add homepage screenshot
    doc.add_heading('Our Website Homepage', level=2)
    doc.add_paragraph('Visit https://hh-asia-tyre.vercel.app to get started:')
    
    screenshot_path = 'c:\\Users\\leonards\\Desktop\\demo4april\\booking_flow_01_homepage.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Figure 1: HH Asia Tyre Homepage')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_page_break()
    
    # ========== SECTION 2: QUICK START ==========
    doc.add_heading('2. Quick Start - Book in 3 Minutes', level=1)
    
    doc.add_heading('For First-Time Users:', level=2)
    
    quick_steps = [
        ('Visit our website', 'Go to: https://hh-asia-tyre.vercel.app'),
        ('Click "Book Now"', 'Located in the top navigation bar or homepage hero section'),
        ('Fill in 4 quick steps', 'Choose branch → Select services → Pick date & time → Enter your info'),
        ('Submit & Wait for Confirmation', 'You\'ll receive a booking reference number!')
    ]
    
    for i, (title, desc) in enumerate(quick_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'Step {i}: {title}')
        run.font.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(desc)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('💡 Tip: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run('Bookmark our website or install our mobile app for quick access!')
    
    doc.add_page_break()
    
    # ========== SECTION 3: CHOOSE BRANCH ==========
    doc.add_heading('3. Step 1: Choose Your Branch', level=1)
    
    doc.add_paragraph(
        'When you click "Book Now", you\'ll see the branch selection screen. '
        'First, choose your region, then select a specific branch.'
    )
    
    doc.add_heading('What You\'ll See:', level=2)
    
    # Add branch selection screenshot
    screenshot_path = 'c:\\Users\\leonards\\Desktop\\demo4april\\booking_flow_02_branch_selection.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Figure 2: Region and Branch Selection')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_heading('Available Regions:', level=2)
    
    regions = [
        ('Manila', 'Shows 5 branches (Alabang, Bicutan, Bacoor, Sucat x2)'),
        ('Laoag', 'Shows 1 branch (Laoag City)')
    ]
    
    for region, desc in regions:
        p = doc.add_paragraph()
        run = p.add_run(f'{region}: ')
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_heading('How to Select a Branch:', level=2)
    
    steps = [
        'Click on your region (Manila or Laoag)',
        'A list of branches will appear',
        'Click on your preferred branch',
        'Review branch details (address, hours, services)',
        'Click "Next: Choose Services"'
    ]
    
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('💡 Tip: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    p.add_run('All 6 branches are now open for online booking!')
    
    doc.add_page_break()
    
    # ========== SECTION 4: VEHICLE DETAILS ==========
    doc.add_heading('4. Step 2: Enter Vehicle Details', level=1)
    
    doc.add_paragraph(
        'Next, enter your vehicle information so our technicians know what they\'ll be working on.'
    )
    
    # Add vehicle screenshot
    screenshot_path = 'c:\\Users\\leonards\\Desktop\\demo4april\\booking_flow_03_vehicle_details.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Figure 3: Vehicle Information Form')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_heading('Required Information:', level=2)
    
    # Create table for vehicle fields
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = ['Field', 'Example']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    vehicle_fields = [
        ('Year', '2024'),
        ('Make', 'Toyota'),
        ('Model', 'Fortuner'),
        ('Plate Number', 'ABC 1234')
    ]
    
    for row_idx, (field, example) in enumerate(vehicle_fields, 1):
        table.rows[row_idx].cells[0].text = field
        table.rows[row_idx].cells[1].text = example
    
    doc.add_page_break()
    
    # ========== SECTION 5: SERVICE SELECTION ==========
    doc.add_heading('5. Step 3: Select Your Services', level=1)
    
    doc.add_paragraph(
        'Choose the services you need. You can select multiple services for your appointment.'
    )
    
    # Add service selection screenshot
    screenshot_path = 'c:\\Users\\leonards\\Desktop\\demo4april\\booking_flow_04_service_selection.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Figure 4: Service Selection')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_heading('Available Services:', level=2)
    
    # Create services table
    services_table = doc.add_table(rows=7, cols=3)
    services_table.style = 'Light Grid Accent 1'
    
    # Headers
    for i, header in enumerate(['Service', 'Description', 'Est. Time']):
        cell = services_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    services = [
        ('Tire Rotation', 'Even tire wear for longer life', '~30 min'),
        ('Oil Change', 'Full oil & filter replacement', '~20 min'),
        ('Brake Inspection', 'Safety check for brake system', '~45 min'),
        ('Battery Replacement', 'Test & replace old batteries', '~30 min'),
        ('Tire Replacement', 'New tires installation', '~45 min'),
        ('General Maintenance', 'Multi-point vehicle check', '~60 min')
    ]
    
    for row_idx, (service, desc, time) in enumerate(services, 1):
        services_table.rows[row_idx].cells[0].text = service
        services_table.rows[row_idx].cells[1].text = desc
        services_table.rows[row_idx].cells[2].text = time
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('⚠️ Note: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    p.add_run('Selecting multiple services may affect your appointment duration.')
    
    doc.add_page_break()
    
    # ========== SECTION 6: DATE & TIME ==========
    doc.add_heading('6. Step 4: Pick Date & Time', level=1)
    
    doc.add_paragraph(
        'Choose your preferred appointment date and time. Available slots are shown in green.'
    )
    
    # Add date/time screenshot
    screenshot_path = 'c:\\Users\\leonards\\Desktop\\demo4april\\booking_flow_05b_date_time_with_slots.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Figure 5: Date and Time Slot Selection')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_heading('Available Options:', level=2)
    
    options = [
        ('Dates:', 'Today up to 6 months in advance'),
        ('Times:', '8:00 AM to 5:00 PM (hourly slots)')
    ]
    
    for label, desc in options:
        p = doc.add_paragraph()
        run = p.add_run(label + ' ')
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_heading('Understanding Time Slot Status:', level=2)
    
    # Status table
    status_table = doc.add_table(rows=4, cols=3)
    status_table.style = 'Light Grid Accent 1'
    
    for i, header in enumerate(['Status', 'Color', 'Meaning']):
        cell = status_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    statuses = [
        ('Open', 'Green', 'Available - Book now!'),
        ('Full', 'Red', 'Already booked - Choose another time'),
        ('Past', 'Gray', 'Time has passed - Cannot select')
    ]
    
    for row_idx, (status, color, meaning) in enumerate(statuses, 1):
        status_table.rows[row_idx].cells[0].text = status
        status_table.rows[row_idx].cells[1].text = color
        status_table.rows[row_idx].cells[2].text = meaning
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('⚠️ IMPORTANT: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run('Only 1 customer can book each time slot. Once booked, the slot shows "Full" immediately.')
    
    doc.add_page_break()
    
    # ========== SECTION 7: CUSTOMER INFO ==========
    doc.add_heading('7. Step 5: Enter Your Information', level=1)
    
    doc.add_paragraph(
        'Provide your contact details so we can send you booking confirmation and updates.'
    )
    
    # Add customer info screenshot
    screenshot_path = 'c:\\Users\\leonards\\Desktop\\demo4april\\booking_flow_06_customer_info.png'
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Figure 6: Customer Information Form')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_heading('Required Information:', level=2)
    
    # Customer info table
    customer_table = doc.add_table(rows=8, cols=2)
    customer_table.style = 'Light Grid Accent 1'
    
    for i, header in enumerate(['Field', 'Example']):
        cell = customer_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    customer_fields = [
        ('Full Name', 'Juan Dela Cruz'),
        ('Email Address', 'juan.delacruz@email.com'),
        ('Phone Number', '0917 123 4567'),
        ('Vehicle Year', '2024'),
        ('Vehicle Make', 'Toyota'),
        ('Vehicle Model', 'Fortuner'),
        ('Plate Number', 'ABC 1234')
    ]
    
    for row_idx, (field, example) in enumerate(customer_fields, 1):
        customer_table.rows[row_idx].cells[0].text = field
        customer_table.rows[row_idx].cells[1].text = example
    
    doc.add_paragraph('')
    doc.add_heading('Validation Rules:', level=2)
    
    rules = [
        'Email must be valid format (name@domain.com)',
        'Phone must be Philippine number (starts with 09 or +63)',
        'Plate number is required'
    ]
    
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== SECTION 8: BRANCH LOCATIONS ==========
    doc.add_heading('8. Our Branch Locations', level=1)
    
    doc.add_paragraph('We have 6 branches ready to serve you:')
    
    doc.add_heading('Metro Manila & Cavite (5 Branches)', level=2)
    
    # Branch locations table
    branch_table = doc.add_table(rows=6, cols=4)
    branch_table.style = 'Light Grid Accent 1'
    
    for i, header in enumerate(['Branch', 'Location', 'Phone', 'Hours']):
        cell = branch_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    branches = [
        ('Alabang', 'Parañaque', '8334-2858', 'Mon-Sat 8AM-6PM'),
        ('Bicutan', 'Taguig', '[Phone]', 'Mon-Sat 8AM-6PM'),
        ('Bacoor', 'Cavite', '(046) 417-8415', 'Mon-Sat 8AM-6PM'),
        ('Sucat (Goodyear)', 'Parañaque', '8828-8050', 'Mon-Sat 8AM-6PM'),
        ('Sucat (GT Radial)', 'Parañaque', '8828-8050', 'Mon-Sat 8AM-6PM')
    ]
    
    for row_idx, (branch, loc, phone, hours) in enumerate(branches, 1):
        branch_table.rows[row_idx].cells[0].text = branch
        branch_table.rows[row_idx].cells[1].text = loc
        branch_table.rows[row_idx].cells[2].text = phone
        branch_table.rows[row_idx].cells[3].text = hours
    
    doc.add_paragraph('')
    doc.add_heading('Ilocos Norte (1 Branch)', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Laoag City - Goodyear Servitek')
    run.font.bold = True
    
    doc.add_paragraph('Address: Bacarra Road, Laoag City, 2900 Ilocos Norte')
    doc.add_paragraph('Phone: (077) 772-3456 | 0917 123 4567')
    doc.add_paragraph('Hours: Mon-Sat 8AM-5PM')
    
    doc.add_page_break()
    
    # ========== SECTION 9: BOOKING STATUS ==========
    doc.add_heading('9. Understanding Booking Status', level=1)
    
    doc.add_paragraph(
        'After submitting your booking, you can track its status. Here\'s what each status means:'
    )
    
    # Status explanation table
    status_exp_table = doc.add_table(rows=7, cols=3)
    status_exp_table.style = 'Light Grid Accent 1'
    
    for i, header in enumerate(['Status', 'Icon', 'What It Means']):
        cell = status_exp_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    booking_statuses = [
        ('Pending', '⏳', 'Waiting for admin approval'),
        ('Approved', '✅', 'Confirmed, bay assigned'),
        ('In Progress', '🔧', 'Service being performed'),
        ('Completed', '✓', 'Service finished'),
        ('Rejected', '❌', 'Booking declined (see reason)'),
        ('Cancelled', '🚫', 'Booking cancelled by you or admin')
    ]
    
    for row_idx, (status, icon, meaning) in enumerate(booking_statuses, 1):
        status_exp_table.rows[row_idx].cells[0].text = status
        status_exp_table.rows[row_idx].cells[1].text = icon
        status_exp_table.rows[row_idx].cells[2].text = meaning
    
    doc.add_page_break()
    
    # ========== SECTION 10: FAQ ==========
    doc.add_heading('10. Frequently Asked Questions', level=1)
    
    faqs = [
        ('How far in advance can I book?',
         'You can book from today up to 6 months in advance.'),
        
        ('Why can\'t I select a time slot?',
         'The slot is either already booked (shows "Full" in red), in the past (shows "Past" in gray), '
         'or outside operating hours.'),
        
        ('Can I book multiple services?',
         'Yes! Select all services you need. The branch will estimate total time.'),
        
        ('How do I know my booking was successful?',
         'You\'ll see a confirmation page with booking reference, receive email confirmation, '
         'and get SMS notification (if enabled).'),
        
        ('Do I need to pay when booking?',
         'No payment is required to book. Pay at the branch after service.'),
        
        ('Can I book same-day appointments?',
         'Yes! If there are available time slots remaining.'),
        
        ('What if I didn\'t receive my confirmation email?',
         'Check your spam/junk folder, wait 5-10 minutes for email delays, verify you entered '
         'the correct email, or contact the branch to confirm booking.')
    ]
    
    for question, answer in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f'Q: {question}')
        run.font.bold = True
        run.font.size = Pt(12)
        
        doc.add_paragraph(f'A: {answer}')
        doc.add_paragraph('')
    
    doc.add_page_break()
    
    # ========== SECTION 11: SUPPORT ==========
    doc.add_heading('11. Need Help?', level=1)
    
    doc.add_heading('Contact Us:', level=2)
    
    contact_info = [
        ('General Inquiries:', ''),
        ('📧 Email:', 'support@hh-asia.com'),
        ('📞 Phone:', '(02) 8123-4567'),
        ('💬 Viber:', '[Viber number]'),
        ('', ''),
        ('Support Hours:', ''),
        ('', 'Monday - Saturday: 8:00 AM - 6:00 PM'),
        ('', 'Sunday: Closed'),
        ('', ''),
        ('🚨 Emergency Roadside Assistance:', ''),
        ('', 'Hotline: 0917-HH-TYRE (0917-448-973)'),
        ('', 'Available 24/7 for emergencies')
    ]
    
    for label, value in contact_info:
        p = doc.add_paragraph()
        if label:
            run = p.add_run(label)
            run.font.bold = True
        if value:
            p.add_run(' ' + value)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Final thank you
    thank_you = doc.add_paragraph()
    thank_you.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = thank_you.add_run('Thank you for choosing HH Asia Tyre! 🚗✨')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    
    doc.add_paragraph('')
    
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final.add_run('Your trusted partner in automotive care')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Save document
    output_path = 'c:\\Users\\leonards\\Desktop\\demo4april\\HH_Asia_Tyre_Customer_Guide.docx'
    doc.save(output_path)
    print(f'✅ Document saved to: {output_path}')
    
    return output_path

if __name__ == '__main__':
    create_customer_guide()
